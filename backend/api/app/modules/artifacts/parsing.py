"""Deterministic, bounded P1 document parsing and chunking."""

from dataclasses import dataclass
from io import BytesIO

import pymupdf
from docx import Document

PARSER_VERSION = "p1.1"
CHUNKER_VERSION = "chars-1200-overlap-150-v1"


class DocumentValidationError(ValueError):
    def __init__(self, code: str, message: str) -> None:
        super().__init__(message)
        self.code = code


@dataclass(frozen=True, slots=True)
class ParsedPage:
    number: int
    text: str
    heading: str | None = None


@dataclass(frozen=True, slots=True)
class ParsedDocument:
    media_type: str
    pages: list[ParsedPage]


@dataclass(frozen=True, slots=True)
class TextChunk:
    index: int
    page_number: int
    text: str
    start_offset: int
    end_offset: int


def normalize_text(value: str) -> str:
    lines = [" ".join(line.split()) for line in value.replace("\x00", "").splitlines()]
    return "\n".join(line for line in lines if line).strip()


def parse_document(
    data: bytes,
    declared_media_type: str,
    *,
    max_pages: int,
    max_characters: int,
) -> ParsedDocument:
    if not data:
        raise DocumentValidationError("artifact_empty", "The uploaded file is empty.")
    if declared_media_type == "application/pdf":
        parsed = parse_pdf(data, max_pages=max_pages)
    elif (
        declared_media_type
        == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ):
        parsed = parse_docx(data, max_pages=max_pages)
    elif declared_media_type == "text/plain":
        parsed = parse_text(data)
    else:
        raise DocumentValidationError("artifact_type_not_allowed", "Unsupported file type.")
    total = sum(len(page.text) for page in parsed.pages)
    if total == 0:
        raise DocumentValidationError("artifact_empty", "No extractable text was found.")
    if total > max_characters:
        raise DocumentValidationError(
            "artifact_text_too_large", "Extracted text exceeds the P1 processing limit."
        )
    return parsed


def parse_pdf(data: bytes, *, max_pages: int) -> ParsedDocument:
    if not data.startswith(b"%PDF"):
        raise DocumentValidationError("artifact_corrupt", "File content is not a valid PDF.")
    try:
        with pymupdf.open(stream=data, filetype="pdf") as document:  # type: ignore[no-untyped-call]
            if document.needs_pass:
                raise DocumentValidationError(
                    "artifact_encrypted", "Encrypted PDFs are not supported in P1."
                )
            if document.page_count > max_pages:
                raise DocumentValidationError(
                    "artifact_page_limit", f"PDF exceeds the {max_pages}-page processing limit."
                )
            pages = [
                ParsedPage(number=index + 1, text=normalize_text(page.get_text("text")))
                for index, page in enumerate(document)
            ]
    except DocumentValidationError:
        raise
    except Exception as exc:
        raise DocumentValidationError("artifact_corrupt", "PDF could not be parsed.") from exc
    return ParsedDocument(media_type="application/pdf", pages=pages)


def parse_docx(data: bytes, *, max_pages: int) -> ParsedDocument:
    if not data.startswith(b"PK"):
        raise DocumentValidationError("artifact_corrupt", "File content is not a valid DOCX.")
    try:
        document = Document(BytesIO(data))
        blocks = [paragraph.text for paragraph in document.paragraphs]
        for table in document.tables:
            blocks.extend(" | ".join(cell.text for cell in row.cells) for row in table.rows)
        text = normalize_text("\n".join(blocks))
    except Exception as exc:
        raise DocumentValidationError("artifact_corrupt", "DOCX could not be parsed.") from exc
    if max_pages < 1:
        raise DocumentValidationError("artifact_page_limit", "Document page limit is invalid.")
    return ParsedDocument(
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        pages=[ParsedPage(number=1, text=text)],
    )


def parse_text(data: bytes) -> ParsedDocument:
    if b"\x00" in data[:4096]:
        raise DocumentValidationError(
            "artifact_corrupt", "Text file appears to contain binary data."
        )
    try:
        text = data.decode("utf-8-sig")
    except UnicodeDecodeError as exc:
        raise DocumentValidationError(
            "artifact_corrupt", "TXT files must use UTF-8 encoding."
        ) from exc
    return ParsedDocument(
        media_type="text/plain", pages=[ParsedPage(number=1, text=normalize_text(text))]
    )


def chunk_document(
    document: ParsedDocument, *, size: int = 1200, overlap: int = 150
) -> list[TextChunk]:
    chunks: list[TextChunk] = []
    index = 0
    for page in document.pages:
        start = 0
        text = page.text
        while start < len(text):
            hard_end = min(len(text), start + size)
            end = hard_end
            if hard_end < len(text):
                boundary = text.rfind(" ", start + size // 2, hard_end)
                if boundary > start:
                    end = boundary
            chunks.append(
                TextChunk(
                    index=index,
                    page_number=page.number,
                    text=text[start:end].strip(),
                    start_offset=start,
                    end_offset=end,
                )
            )
            index += 1
            if end >= len(text):
                break
            start = max(start + 1, end - overlap)
    return [chunk for chunk in chunks if chunk.text]
